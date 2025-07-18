from html2docx import html2docx

def convert_html_to_docx_content(html_path: Path) -> Document:
    document = Document()
    html = html_path.read_text(encoding="utf-8")
    html2docx(html, document)
    return document


from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from pathlib import Path
from docx import Document
import uuid
from bs4 import BeautifulSoup

app = FastAPI()

class DocumentRequest(BaseModel):
    document_name: str
    class_id: str
    type: str

DOCUMENTS_DIR = Path("common")
OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)

def convert_html_to_docx_content(html_path: Path) -> Document:
    from html2text import html2text
    from docx import Document

    document = Document()
    html = html_path.read_text(encoding="utf-8")

    # Optional: clean and parse with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Basic conversion: convert text content
    markdown_text = html2text(str(soup))

    for line in markdown_text.splitlines():
        document.add_paragraph(line)

    return document

@app.post("/process-document/")
async def process_document(request: DocumentRequest):
    html_path = DOCUMENTS_DIR / request.document_name

    if not html_path.exists():
        raise HTTPException(status_code=404, detail="HTML file not found.")

    try:
        # Read the HTML file content
        html_content = html_path.read_text(encoding="utf-8")

        # Split the document at <results> placeholders
        sections = html_content.split("<results>")
        if len(sections) == 1:
            raise HTTPException(status_code=400, detail="No <results> placeholder found.")

        # Convert HTML to a DOCX Document
        result_docx = convert_html_to_docx_content(html_path)

        # Build final Word document
        final_doc = Document()

        for i, section in enumerate(sections):
            # Add HTML section (before or between results)
            final_doc.add_paragraph(BeautifulSoup(section, "html.parser").get_text())

            # Add result_docx pages if this is not the last section
            if i < len(sections) - 1:
                # Insert a page break before inserting result content
                final_doc.add_page_break()
                for para in result_docx.paragraphs:
                    final_doc.add_paragraph(para.text)
                final_doc.add_page_break()

        # Generate final filename
        base_name = Path(request.document_name).stem
        safe_class_id = request.class_id.replace(" ", "_")
        safe_type = request.type.replace(" ", "_")
        output_filename = f"{base_name}_{safe_class_id}_{safe_type}.docx"
        output_path = OUTPUT_DIR / output_filename

        # Save the final DOCX
        final_doc.save(str(output_path))

        return {
            "message": "HTML converted, processed, and saved as DOCX.",
            "output_file": str(output_path)
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

