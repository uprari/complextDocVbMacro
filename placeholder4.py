from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from html2docx import html2docx
from bs4 import BeautifulSoup

app = FastAPI()

DOCUMENTS_DIR = Path("common")
OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)

class DocumentRequest(BaseModel):
    document_name: str      # DOCX file
    class_id: str
    type: str
    html_file: str          # HTML file used to insert content

def insert_page_break(paragraph):
    """Inserts a hard page break after a paragraph."""
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

@app.post("/process-docx/")
async def process_docx(request: DocumentRequest):
    try:
        docx_path = DOCUMENTS_DIR / request.document_name
        html_path = DOCUMENTS_DIR / request.html_file

        if not docx_path.exists():
            raise HTTPException(status_code=404, detail="DOCX file not found.")
        if not html_path.exists():
            raise HTTPException(status_code=404, detail="HTML file not found.")

        # Load DOCX file
        doc = Document(docx_path)

        # Find and replace placeholder <results>
        placeholder_found = False
        for i, paragraph in enumerate(doc.paragraphs):
            if "<results>" in paragraph.text:
                placeholder_found = True
                # Remove the placeholder
                paragraph.text = paragraph.text.replace("<results>", "")
                # Insert page break after this paragraph
                insert_page_break(paragraph)

                # Now convert HTML to a new DOCX section
                html_doc = Document()
                html = html_path.read_text(encoding="utf-8")
                html2docx(html, html_doc)

                # Insert converted HTML content into original doc
                for para in html_doc.paragraphs:
                    new_para = doc.add_paragraph()
                    new_para.style = para.style
                    new_para.add_run(para.text)

                break  # Only one placeholder expected

        if not placeholder_found:
            raise HTTPException(status_code=400, detail="No <results> placeholder found in DOCX.")

        # Generate output file name
        base_name = Path(request.document_name).stem
        safe_class_id = request.class_id.replace(" ", "_")
        safe_type = request.type.replace(" ", "_")
        output_filename = f"{base_name}_{safe_class_id}_{safe_type}.docx"
        output_path = OUTPUT_DIR / output_filename

        # Save the final document
        doc.save(str(output_path))

        return {
            "message": "Document processed and saved.",
            "output_file": str(output_path)
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

