import re
import json
import base64
from abc import ABC, abstractmethod
from typing import Dict, Tuple
from docx import Document
from docx.shared import Inches
from io import BytesIO


class AbstractReplacer(ABC):
    @abstractmethod
    def replace(self, document, paragraph, parent: str, key: str):
        pass


class MetricReplacer(AbstractReplacer):
    def __init__(self, data: Dict[Tuple[str, str], Dict]):
        self.data = data

    def replace(self, document, paragraph, parent: str, key: str):
        value = self.data.get((parent, key), {}).get("value", [])
        if isinstance(value, list):
            idx = paragraph._p.getparent().index(paragraph._p)
            parent_element = paragraph._p.getparent()
            for i, line in enumerate(value):
                new_para = document.paragraphs[idx + i]
                new_para.text = line
            paragraph.text = ""
        else:
            paragraph.text = str(value)


class TableReplacer(AbstractReplacer):
    def __init__(self, data: Dict[Tuple[str, str], Dict]):
        self.data = data

    def replace(self, document, paragraph, parent: str, key: str):
        value = self.data.get((parent, key), {}).get("value", {})
        columns = value.get("column", [])
        data_rows = value.get("data", [])

        idx = paragraph._p.getparent().index(paragraph._p)
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = col

        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)

        # Insert table and remove original paragraph
        paragraph.text = ""
        paragraph._p.addnext(table._tbl)


class FigureReplacer(AbstractReplacer):
    def __init__(self, data: Dict[Tuple[str, str], Dict]):
        self.data = data

    def replace(self, document, paragraph, parent: str, key: str):
        encoded_img = self.data.get((parent, key), {}).get("value")
        if encoded_img:
            image_data = base64.b64decode(encoded_img)
            image_stream = BytesIO(image_data)
            run = paragraph.clear().add_run()
            run.add_picture(image_stream, width=Inches(4))
        else:
            paragraph.text = "[Missing Image]"


class ReplacerFactory:
    def __init__(self, metric_data, table_data, figure_data):
        self._replacers = {
            "metric": MetricReplacer(metric_data),
            "table": TableReplacer(table_data),
            "figure": FigureReplacer(figure_data)
        }

    def get_replacer(self, type_key: str) -> AbstractReplacer:
        return self._replacers.get(type_key)


class PatternParser:
    pattern = re.compile(r"<(?P<parent>[^.]+)\.(?P<type>figure|table|metric)\.(?P<key>[^>]+)>")

    @staticmethod
    def parse(text: str):
        return PatternParser.pattern.finditer(text)


class DocxReplacer:
    def __init__(self, docx_path: str, replacer_factory: ReplacerFactory):
        self.doc = Document(docx_path)
        self.factory = replacer_factory

    def process(self):
        for para in self.doc.paragraphs:
            matches = list(PatternParser.parse(para.text))
            for match in matches:
                full_match = match.group(0)
                parent = match.group("parent")
                type_key = match.group("type")
                key = match.group("key")
                replacer = self.factory.get_replacer(type_key)
                if replacer:
                    replacer.replace(self.doc, para, parent, key)

    def save(self, output_path: str):
        self.doc.save(output_path)


def load_nested_result(json_path: str):
    with open(json_path, "r", encoding="utf-8") as f:
        raw = json.load(f)

    metric_data = {}
    table_data = {}
    figure_data = {}

    for entry in raw.get("_result", []):
        for parent, content in entry.items():
            for type_, items in content.items():
                for key, meta in items.items():
                    if type_ == "metric":
                        metric_data[(parent, key)] = meta
                    elif type_ == "table":
                        table_data[(parent, key)] = meta
                    elif type_ == "figure":
                        figure_data[(parent, key)] = meta

    return metric_data, table_data, figure_data


# ---------- Usage ----------
if __name__ == "__main__":
    json_path = "result.json"
    input_docx = "input.docx"
    output_docx = "output.docx"

    metric, table, figure = load_nested_result(json_path)
    factory = ReplacerFactory(metric, table, figure)

    doc_proc = DocxReplacer(input_docx, factory)
    doc_proc.process()
    doc_proc.save(output_docx)

