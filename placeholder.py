from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from docx import Document
import os

app = FastAPI()

COMMONS_FOLDER = "commons"

class DocumentRequest(BaseModel):
    name: str
    school_id: str
    type: str
    result: str = "Sample Result"  # Optional default; could be dynamic


@app.post("/process-document/")
def process_document(data: DocumentRequest):
    input_filename = os.path.join(COMMONS_FOLDER, f"{data.name}.docx")
    
    if not os.path.isfile(input_filename):
        raise HTTPException(status_code=404, detail="Document not found.")

    # Load document
    doc = Document(input_filename)

    # Replace placeholder with result in paragraphs
    for para in doc.paragraphs:
        if "<REPLACE_RESULT>" in para.text:
            para.text = para.text.replace("<REPLACE_RESULT>", data.result)

    # Also check in tables (if needed)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "<REPLACE_RESULT>" in cell.text:
                    cell.text = cell.text.replace("<REPLACE_RESULT>", data.result)

    # Create new file name
    new_filename = f"{data.name}_{data.school_id}_{data.type}.docx"
    new_filepath = os.path.join(COMMONS_FOLDER, new_filename)

    # Save updated document
    doc.save(new_filepath)

    return {
        "message": "Document processed and saved successfully.",
        "output_file": new_filename
    }
